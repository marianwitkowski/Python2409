
# Tworzenie dokumentu DOCX
from docx import Document

document : Document = Document()
#print(document)
document.add_heading("Nagłówek")
document.add_paragraph("Tu jest jakiś paragraf 1")
par2 = document.add_paragraph("Tu jest jakiś paragraf 2")
par2.insert_paragraph_before("Tekst przed paragrafer nr 2")

par = document.add_paragraph()
par.add_run("Treść zwykła ")
par.add_run("Treść wyboldowana ").bold = True
par.add_run("Treść zwykła ")

document.add_page_break()

# Rysowanie tablicy
data = [
    (1, "Komputer Dell", 3499),
    (2, "Redmi 10 Note", 1699),
    (3, "Dysk 4TB", 699),
]
# dodaj tabele
table = document.add_table(1, 3)
header = table.rows[0].cells # zwraca na liste komorek w ramach 1-go wiersza
header[0].text = "Lp."
header[1].text = "Produkt"
header[2].text = "Cena"

for item in data:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = str(item[2])

document.save("test.docx")
